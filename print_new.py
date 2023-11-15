from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import RGBColor


# Function to add a row to a table with specific font settings
def add_row_to_table(table, row_data, font_size, font_color, bold=False, underline=False):
    row = table.add_row().cells
    for idx, data in enumerate(row_data):
        run = row[idx].paragraphs[0].add_run(data)
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(*font_color)
        run.font.bold = bold
        run.font.underline = underline

# Initialize a Document
doc = Document()

# Add a table with 3 columns
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'

# Define colors
maroon = (128, 0, 0)
navy = (0, 0, 128)
black = (0, 0, 0)

# Add header row
add_row_to_table(table, ["Class Name", "Teachers", "Week of"], 14, maroon, bold=True, underline=True)

# Add user input row (replace 'class_name', 'teachers', 'week_of' with actual user input)
add_row_to_table(table, ['class_name', 'teachers', 'week_of'], 12, navy)

# Add theme rows
doc.add_paragraph('Theme', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
theme_run = doc.paragraphs[-1].add_run()
theme_run.font.size = Pt(14)
theme_run.font.color.rgb = RGBColor(*maroon)
theme_run.font.underline = True

# Add user input for theme (replace 'user_theme' with actual user input)
doc.add_paragraph('user_theme', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
theme_run = doc.paragraphs[-1].add_run()
theme_run.font.size = Pt(12)
theme_run.font.color.rgb = RGBColor(*navy)

# Add a blank line
doc.add_paragraph()

# Add activities and skills (replace these lists with actual user input)
days = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"]
activities = ["Activity for Monday", "Activity for Tuesday", "Activity for Wednesday", "Activity for Thursday", "Activity for Friday"]
skills = ["Skills for Monday", "Skills for Tuesday", "Skills for Wednesday", "Skills for Thursday", "Skills for Friday"]

for day, activity, skill in zip(days, activities, skills):
    # Add day
    day_para = doc.add_paragraph()
    day_run = day_para.add_run(day + ": ")
    day_run.font.size = Pt(14)
    day_run.font.color.rgb = RGBColor(*maroon)
    day_run.font.underline = True

    # Add activity
    activity_run = day_para.add_run(activity)
    activity_run.font.size = Pt(14)
    activity_run.font.color.rgb = RGBColor(*black)

    # Add skills
    skills_para = doc.add_paragraph("Skills: ")
    skills_run = skills_para.add_run(skill)
    skills_run.font.size = Pt(14)
    skills_run.font.color.rgb = RGBColor(*navy)

# Save the document
doc.save('class_schedule.docx')
