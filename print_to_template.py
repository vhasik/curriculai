from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


# Function to replace text in a paragraph while keeping the formatting
def replace_paragraph_text(paragraph, new_text):
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ''


def set_cell_text(cell, text):
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def find_and_duplicate_section(doc, placeholder, content_list):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # Found the placeholder, now replicate and replace for each item in content_list
            style = paragraph.style
            for content in content_list:
                # Create a new paragraph with the same style
                new_para = doc.add_paragraph(style=style)
                new_para.add_run(content)
            # Remove the placeholder paragraph
            p = paragraph._element
            p.getparent().remove(p)
            break


# Load your template
template_path = os.path.join("templates", "curriculum_template.docx")
doc = Document(template_path)

# Replace placeholders for activities and skills
week_of = "1/1/23 - 1/7/23"
class_name = "My Class"
teachers = "Teacher 1, Teacher 2"
theme = "My Theme"
days = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"]
activities = ["Activity for Monday", "Activity for Tuesday", "Activity for Wednesday", "Thursday activity", "Friday activity"]
skills = ["Skills for Monday", "Skills for Tuesday", "Skills for Wednesday", "Skills for Thursday", "Skills for Friday"]

# Replace placeholders in specific table cells
table = doc.tables[0]
set_cell_text(table.cell(1, 0), 'New Class Name')  # Second row, first column
set_cell_text(table.cell(1, 1), 'New Teachers')    # Second row, second column
set_cell_text(table.cell(1, 2), 'New Week Of')     # Second row, third column

# Ensure the paragraphs are centered
for cell in table.rows[1].cells:  # Adjust the row index as needed
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Replace placeholders basd on keyword
for paragraph in doc.paragraphs:
    if 'THEME_PLACEHOLDER' in paragraph.text:
        replace_paragraph_text(paragraph, 'Actual Theme')
    # Add more elif conditions for other placeholders

# Replace placeholder activities and skills
# Define colors
maroon = (128, 0, 0)
green = (0, 69, 0)
black = (0, 0, 0)
navy = (0, 0, 69)

for day, activity, skill in zip(days, activities, skills):
    # Create a new paragraph for each day
    day_para = doc.add_paragraph()

    # Add day and format it
    day_run = day_para.add_run(day + ": ")
    day_run.font.size = Pt(18)
    day_run.font.color.rgb = RGBColor(*maroon)
    day_run.font.underline = True
    day_run.font.bold = True

    # Add activity and format it
    activity_run = day_para.add_run(activity + "\n")
    activity_run.font.size = Pt(16)
    activity_run.font.color.rgb = RGBColor(*green)
    activity_run.font.bold = True

    # Add skills label and format it
    skills_run = day_para.add_run("Skills: ")
    skills_run.font.size = Pt(16)
    skills_run.font.color.rgb = RGBColor(*navy)
    skills_run.font.bold = True

    # Add skills and format it
    skills_run = day_para.add_run(skill)
    skills_run.font.size = Pt(16)
    skills_run.font.color.rgb = RGBColor(*navy)
    skills_run.font.bold = True


# Save the document as a new file
doc.save('new_class_schedule.docx')
