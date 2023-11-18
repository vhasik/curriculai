from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta
import os


# Function to replace text in a paragraph while keeping the formatting
def replace_paragraph_text(paragraph, new_text):
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ''


def set_cell_text(cell, text):
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def find_and_duplicate_section(doc, placeholder, content_list):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # Found the placeholder, now replicate and replace for each item in content_list
            style = paragraph.style
            for content in content_list:
                # Create a new paragraph with the same style
                new_para = doc.add_paragraph(style=style)
                new_para.add_run(content)
            # Remove the placeholder paragraph
            p = paragraph._element
            p.getparent().remove(p)
            break


def format_week_of(input_date):
    if not input_date:
        return "Unknown Date Range"

    try:
        print(f"Input date: {input_date}")
        start_date_string = input_date.strftime('%m/%d/%y')
    except ValueError as e:
        print(f"Error parsing date: {e}")
        return "Invalid Date Format"

    print(f"Start date: {start_date_string}")
    end_date = input_date + timedelta(days=5)
    end_date_string = end_date.strftime('%m/%d/%y')

    print(f"End date: {end_date_string}")

    return f"{start_date_string} - {end_date_string}"


def get_test_data():
    # Replace this with your own data
    class_name = "My Class"
    teachers = "Teacher 1, Teacher 2"
    week_of = "1/1/23 - 1/7/23"
    theme = "My Theme"
    days = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"]
    activities = ["Activity for Monday", "Activity for Tuesday", "Activity for Wednesday", "Thursday activity", "Friday activity"]
    skills = ["Skills for Monday", "Skills for Tuesday", "Skills for Wednesday", "Skills for Thursday", "Skills for Friday"]
    return class_name, teachers, week_of, theme, days, activities, skills


def generate_doc(class_name, teachers, input_date, theme, days, activities, skills):
    # Format the date
    print("Formatting date...")
    new_date = datetime.strptime(input_date, '%Y-%m-%d')
    date = new_date.strftime('%y-%m-%d')
    week_of = format_week_of(new_date)

    # Load your template
    print("Loading template...")
    template_path = os.path.join("templates", "curriculum_template.docx")
    doc = Document(template_path)

    print("Creating document...")

    print("Days: ", days)
    print("Activities: ", activities)
    print("Skills: ", skills)

    # Replace placeholders in specific table cells
    table = doc.tables[0]
    set_cell_text(table.cell(1, 0), class_name)  # Second row, first column
    set_cell_text(table.cell(1, 1), teachers)    # Second row, second column
    set_cell_text(table.cell(1, 2), week_of)     # Second row, third column

    # Ensure the paragraphs are centered
    for cell in table.rows[1].cells:  # Adjust the row index as needed
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Replace placeholders basd on keyword
    for paragraph in doc.paragraphs:
        if 'THEME_PLACEHOLDER' in paragraph.text:
            replace_paragraph_text(paragraph, theme)
        # Add more elif conditions for other placeholders

    # Replace placeholder activities and skills
    # Define colors
    maroon = (128, 0, 0)
    green = (0, 69, 0)
    black = (0, 0, 0)
    navy = (0, 0, 69)

    for day, activity, skill in zip(days, activities, skills):
        # Create a new paragraph for each day
        day_para = doc.add_paragraph()

        # Add day and format it
        day_run = day_para.add_run(day + ": ")
        day_run.font.size = Pt(18)
        day_run.font.color.rgb = RGBColor(*maroon)
        day_run.font.underline = True
        day_run.font.bold = True

        # Add activity and format it
        activity_run = day_para.add_run(activity + "\n")
        activity_run.font.size = Pt(16)
        activity_run.font.color.rgb = RGBColor(*green)
        activity_run.font.bold = True
        day_para.paragraph_format.space_after = Pt(12)

        # Add skills label and format it
        skills_run = day_para.add_run("Skills: ")
        skills_run.font.size = Pt(16)
        skills_run.font.color.rgb = RGBColor(*navy)
        skills_run.font.bold = True

        # Add skills and format it
        skills_run = day_para.add_run(skill)
        skills_run.font.size = Pt(16)
        skills_run.font.color.rgb = RGBColor(*navy)
        skills_run.font.bold = True
        day_para.paragraph_format.space_before = Pt(6)
        day_para.paragraph_format.space_after = Pt(12)

    print("Saving document...")
    # Save the document as a new file
    new_doc_path = os.path.join("assets", f"curriculum_{date}.docx")
    doc.save(new_doc_path)

    return new_doc_path
